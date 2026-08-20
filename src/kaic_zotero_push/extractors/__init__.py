"""Safe local extraction for supported document formats."""

from __future__ import annotations

import csv
import hashlib
import io
import re
from typing import TYPE_CHECKING

import pymupdf
from docx import Document
from openpyxl import load_workbook

from kaic_zotero_push.errors import InputError
from kaic_zotero_push.extractors.sections import is_reference_heading, reference_paragraphs
from kaic_zotero_push.models import (
    ExtractedDocument,
    SourceCandidate,
    StructuredReference,
)

if TYPE_CHECKING:
    from collections.abc import Callable
    from pathlib import Path

_MAX_INPUT_BYTES = 100 * 1024 * 1024
_REFERENCE_START = re.compile(r"^\s*(?:\[\d+\]|\d+[.)]|[-*+])\s+")
_COLUMN_ALIASES = {
    "title": "title",
    "제목": "title",
    "author": "author",
    "authors": "author",
    "저자": "author",
    "year": "year",
    "date": "year",
    "연도": "year",
    "journal": "container_title",
    "publicationtitle": "container_title",
    "학술지": "container_title",
    "volume": "volume",
    "권": "volume",
    "issue": "issue",
    "호": "issue",
    "pages": "pages",
    "쪽": "pages",
    "publisher": "publisher",
    "출판사": "publisher",
    "doi": "doi",
    "pmid": "pmid",
    "isbn": "isbn",
    "url": "url",
    "itemtype": "item_type",
    "type": "item_type",
}


def _segment_lines(lines: list[str], *, locator_name: str) -> list[SourceCandidate]:
    candidates: list[SourceCandidate] = []
    buffer: list[str] = []
    start_line = 1
    for line_number, line in enumerate(lines, start=1):
        stripped = line.strip()
        starts_reference = bool(_REFERENCE_START.match(stripped))
        if buffer and (starts_reference or not stripped):
            candidates.append(
                SourceCandidate(
                    source_index=len(candidates) + 1,
                    source_locator=f"{locator_name}={start_line}",
                    raw_text=" ".join(buffer),
                )
            )
            buffer = []
        if stripped:
            if not buffer:
                start_line = line_number
            buffer.append(stripped)
    if buffer:
        candidates.append(
            SourceCandidate(
                source_index=len(candidates) + 1,
                source_locator=f"{locator_name}={start_line}",
                raw_text=" ".join(buffer),
            )
        )
    return candidates


def _structured(values: dict[str, str]) -> StructuredReference:
    mapped = {
        target: value.strip()
        for name, value in values.items()
        if (target := _COLUMN_ALIASES.get(name.casefold().replace(" ", ""))) and value.strip()
    }
    return StructuredReference.model_validate(mapped)


def _extract_text(path: Path) -> list[SourceCandidate]:
    try:
        text = path.read_text(encoding="utf-8-sig")
    except UnicodeDecodeError as error:
        raise InputError(code="EXTRACTION_FAILED", detail="Text input must be UTF-8.") from error
    if "\x00" in text:
        raise InputError(code="INPUT_UNSUPPORTED", detail="Binary content is not a text document.")
    lines = [line for line in text.splitlines() if not is_reference_heading(line)]
    return _segment_lines(lines, locator_name="line")


def _extract_csv(path: Path) -> list[SourceCandidate]:
    try:
        text = path.read_text(encoding="utf-8-sig")
    except UnicodeDecodeError as error:
        raise InputError(code="EXTRACTION_FAILED", detail="CSV input must be UTF-8.") from error
    reader = csv.DictReader(io.StringIO(text))
    if reader.fieldnames is None:
        raise InputError(code="EXTRACTION_FAILED", detail="CSV requires a header row.")
    candidates: list[SourceCandidate] = []
    for row_number, row in enumerate(reader, start=2):
        values = {str(name): value or "" for name, value in row.items() if name is not None}
        structured = _structured(values)
        raw = " | ".join(value for value in values.values() if value.strip())
        if raw:
            candidates.append(
                SourceCandidate(
                    source_index=len(candidates) + 1,
                    source_locator=f"row={row_number}",
                    raw_text=raw,
                    structured=structured,
                )
            )
    return candidates


def _extract_xlsx(path: Path) -> list[SourceCandidate]:
    workbook = load_workbook(path, read_only=True, data_only=True)
    try:
        candidates: list[SourceCandidate] = []
        for sheet in workbook.worksheets:
            rows = sheet.iter_rows(values_only=True)
            headers = next(rows, None)
            if headers is None:
                continue
            names = [str(value or "") for value in headers]
            for row_number, row in enumerate(rows, start=2):
                values = {
                    name: str(value)
                    for name, value in zip(names, row, strict=False)
                    if name and value is not None
                }
                raw = " | ".join(values.values())
                if raw:
                    candidates.append(
                        SourceCandidate(
                            source_index=len(candidates) + 1,
                            source_locator=f"sheet={sheet.title};row={row_number}",
                            raw_text=raw,
                            structured=_structured(values),
                        )
                    )
        return candidates
    finally:
        workbook.close()


def _extract_docx(path: Path) -> list[SourceCandidate]:
    document = Document(str(path))
    section_confirmed, paragraphs = reference_paragraphs(document)
    if section_confirmed:
        return [
            SourceCandidate(
                source_index=index,
                source_locator=f"paragraph={paragraph_index}",
                raw_text=text,
            )
            for index, (paragraph_index, text) in enumerate(paragraphs, start=1)
        ]
    blocks: list[tuple[str, str]] = []
    for paragraph_index, paragraph in enumerate(document.paragraphs, start=1):
        if paragraph.text.strip():
            blocks.append((f"paragraph={paragraph_index}", paragraph.text.strip()))
    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if text:
                blocks.append((f"table={table_index};row={row_index}", text))
    return [
        SourceCandidate(
            source_index=index,
            source_locator=locator,
            raw_text=text,
            section_confirmed=False,
        )
        for index, (locator, text) in enumerate(blocks, start=1)
    ]


def _extract_pdf(path: Path) -> list[SourceCandidate]:
    with pymupdf.open(str(path)) as document:
        if document.needs_pass:
            raise InputError(code="INPUT_ENCRYPTED", detail="Unlock the PDF before importing.")
        blocks: list[SourceCandidate] = []
        for page_number, page in enumerate(document, start=1):
            for block_number, block in enumerate(page.get_text("blocks"), start=1):
                text = str(block[4]).strip()
                if text:
                    blocks.append(
                        SourceCandidate(
                            source_index=len(blocks) + 1,
                            source_locator=f"page={page_number};block={block_number}",
                            raw_text=" ".join(text.splitlines()),
                        )
                    )
        if not blocks:
            raise InputError(
                code="EXTRACTION_FAILED",
                detail="PDF has no text layer; scanned PDFs are not supported in v0.2.",
            )
        return blocks


_EXTRACTORS: dict[str, Callable[[Path], list[SourceCandidate]]] = {
    ".txt": _extract_text,
    ".md": _extract_text,
    ".csv": _extract_csv,
    ".xlsx": _extract_xlsx,
    ".docx": _extract_docx,
    ".pdf": _extract_pdf,
}


def extract_document(path: Path) -> ExtractedDocument:
    """Validate, hash, and locally extract one supported document."""
    if not path.is_file():
        raise InputError(code="INPUT_NOT_FOUND", detail=f"File not found: {path}")
    size = path.stat().st_size
    if size > _MAX_INPUT_BYTES:
        raise InputError(code="INPUT_TOO_LARGE", detail="Input exceeds the 100 MiB limit.")
    extractor = _EXTRACTORS.get(path.suffix.casefold())
    if extractor is None:
        raise InputError(
            code="INPUT_UNSUPPORTED",
            detail="Supported formats: .docx, .xlsx, .csv, .md, .txt, text-based .pdf.",
        )
    content = path.read_bytes()
    candidates = extractor(path)
    if not candidates:
        raise InputError(code="EXTRACTION_FAILED", detail="No reference candidates were found.")
    return ExtractedDocument(
        input_path=str(path),
        file_sha256=hashlib.sha256(content).hexdigest(),
        candidates=candidates,
    )
