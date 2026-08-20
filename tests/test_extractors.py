import csv
from pathlib import Path

from docx import Document
from openpyxl import Workbook

from kaic_zotero_push.extractors import extract_document


def test_extract_document_when_text_has_numbered_references(tmp_path: Path) -> None:
    # Given
    path = tmp_path / "references.txt"
    content = """1. Smith, J. (2025). First title. Journal.
2. Lee, M. (2024). Second title. Journal.
"""
    _ = path.write_text(content, encoding="utf-8")

    # When
    extracted = extract_document(path)

    # Then
    assert len(extracted.candidates) == 2
    assert extracted.candidates[0].source_locator == "line=1"


def test_extract_document_when_csv_has_structured_columns(tmp_path: Path) -> None:
    # Given
    path = tmp_path / "references.csv"
    with path.open("w", encoding="utf-8", newline="") as stream:
        writer = csv.DictWriter(stream, fieldnames=["title", "author", "year", "doi"])
        writer.writeheader()
        writer.writerow(
            {
                "title": "A structured title",
                "author": "Smith, John",
                "year": "2025",
                "doi": "10.1000/csv",
            }
        )

    # When
    extracted = extract_document(path)

    # Then
    assert extracted.candidates[0].structured is not None
    assert extracted.candidates[0].structured.title == "A structured title"
    assert extracted.candidates[0].source_locator == "row=2"


def test_extract_document_when_xlsx_has_structured_columns(tmp_path: Path) -> None:
    # Given
    path = tmp_path / "references.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    assert sheet is not None
    sheet.append(["title", "author", "year", "doi"])
    sheet.append(["Spreadsheet title", "Kim, Hana", "2024", "10.1000/xlsx"])
    workbook.save(path)

    # When
    extracted = extract_document(path)

    # Then
    assert extracted.candidates[0].structured is not None
    assert extracted.candidates[0].structured.doi == "10.1000/xlsx"
    assert extracted.candidates[0].source_locator.endswith("row=2")


def test_extract_document_when_docx_has_reference_section_uses_only_that_section(
    tmp_path: Path,
) -> None:
    # Given
    path = tmp_path / "article.docx"
    document = Document()
    _ = document.add_heading("Introduction", level=1)
    _ = document.add_paragraph("This body paragraph cites prior work (Smith, 2024).")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Body table note that is not a reference."
    _ = document.add_heading("References", level=1)
    _ = document.add_paragraph("1. Smith, J. (2025). First title. Journal of Testing, 4(2), 10-20.")
    _ = document.add_paragraph("2. Lee, M. (2024). Second title. Another Journal, 3(1), 21-30.")
    _ = document.add_heading("Appendix", level=1)
    _ = document.add_paragraph("Supplementary body text.")
    document.save(str(path))

    # When
    extracted = extract_document(path)

    # Then
    assert [candidate.raw_text for candidate in extracted.candidates] == [
        "1. Smith, J. (2025). First title. Journal of Testing, 4(2), 10-20.",
        "2. Lee, M. (2024). Second title. Another Journal, 3(1), 21-30.",
    ]
    assert all(candidate.section_confirmed for candidate in extracted.candidates)


def test_extract_document_when_docx_has_no_reference_heading_marks_boundary_unconfirmed(
    tmp_path: Path,
) -> None:
    # Given
    path = tmp_path / "reference-list.docx"
    document = Document()
    _ = document.add_paragraph("1. Smith, J. (2025). First title. Journal.")
    document.save(str(path))

    # When
    extracted = extract_document(path)

    # Then
    assert len(extracted.candidates) == 1
    assert extracted.candidates[0].section_confirmed is False


def test_extract_document_when_text_has_reference_heading_excludes_heading(
    tmp_path: Path,
) -> None:
    # Given
    path = tmp_path / "references.txt"
    _ = path.write_text(
        """References
1. Smith, J. (2025). First title. Journal.
2. Lee, M. (2024). Second title. Journal.
""",
        encoding="utf-8",
    )

    # When
    extracted = extract_document(path)

    # Then
    assert len(extracted.candidates) == 2
    assert all(candidate.raw_text != "References" for candidate in extracted.candidates)
