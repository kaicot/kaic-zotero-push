from pathlib import Path

import pytest
from docx import Document

from kaic_zotero_push.extractors import extract_document


def test_extract_document_when_normal_style_supplementary_tables_follow_references(
    tmp_path: Path,
) -> None:
    # Given
    path = tmp_path / "article.docx"
    document = Document()
    _ = document.add_paragraph("Article body.")
    _ = document.add_heading("References", level=1)
    for index in range(1, 28):
        _ = document.add_paragraph(f"{index}. Smith, J. (2025). Reference {index}. Journal.")
    _ = document.add_paragraph("Table S1. Sensitivity analyses for the main models.")
    _ = document.add_paragraph("Values are odds ratios. CI, confidence interval.")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Supplementary table values."
    _ = document.add_paragraph("Table S2. PHQ-8 severity-band sensitivity analysis.")
    _ = document.add_paragraph("Models adjusted for age and sex.")
    document.save(str(path))

    # When
    extracted = extract_document(path)

    # Then
    assert len(extracted.candidates) == 27
    assert all(candidate.section_confirmed for candidate in extracted.candidates)
    assert all(
        candidate.raw_text.startswith(f"{index}. ")
        for index, candidate in enumerate(extracted.candidates, start=1)
    )
    assert all(
        candidate.source_locator not in {"paragraph=30", "paragraph=31", "paragraph=32"}
        for candidate in extracted.candidates
    )


@pytest.mark.parametrize(
    "terminator",
    [
        "Table S1. Sensitivity analyses.",
        "Table S2. Secondary analyses.",
        "Table 1. Participant characteristics.",
        "Supplementary",
        "Supplementary Table S3",
        "Supporting Information",
        "Appendix",
        "Acknowledgments",
        "Figure 1. Study flow.",
    ],
)
def test_extract_document_when_normal_style_terminator_follows_references(
    tmp_path: Path,
    terminator: str,
) -> None:
    # Given
    path = tmp_path / "article.docx"
    document = Document()
    _ = document.add_heading("References", level=1)
    _ = document.add_paragraph("Smith, J. (2025). First title. Journal.")
    _ = document.add_paragraph(terminator)
    _ = document.add_paragraph("Description or footnote that is not a reference.")
    document.save(str(path))

    # When
    extracted = extract_document(path)

    # Then
    assert [candidate.raw_text for candidate in extracted.candidates] == [
        "Smith, J. (2025). First title. Journal."
    ]
    assert extracted.candidates[0].section_confirmed is True


def test_extract_document_when_reference_section_has_no_terminator_requires_review(
    tmp_path: Path,
) -> None:
    # Given
    path = tmp_path / "article.docx"
    document = Document()
    _ = document.add_paragraph("Article body.")
    _ = document.add_heading("References", level=1)
    _ = document.add_paragraph("1. Smith, J. (2025). First title. Journal.")
    _ = document.add_paragraph("2. Lee, M. (2024). Second title. Journal.")
    document.save(str(path))

    # When
    extracted = extract_document(path)

    # Then
    assert len(extracted.candidates) == 2
    assert all(candidate.section_confirmed is False for candidate in extracted.candidates)


def test_extract_document_when_references_are_unnumbered_keeps_them_until_terminator(
    tmp_path: Path,
) -> None:
    # Given
    path = tmp_path / "article.docx"
    document = Document()
    _ = document.add_heading("Bibliography", level=1)
    _ = document.add_paragraph("Smith, J. (2025). First title. Journal.")
    _ = document.add_paragraph("Lee, M. (2024). Second title. Journal.")
    _ = document.add_paragraph("Appendix")
    _ = document.add_paragraph("Supplementary body text.")
    document.save(str(path))

    # When
    extracted = extract_document(path)

    # Then
    assert [candidate.raw_text for candidate in extracted.candidates] == [
        "Smith, J. (2025). First title. Journal.",
        "Lee, M. (2024). Second title. Journal.",
    ]
